"""Test smoke degli extractor: ogni tipo produce un ExtractionResult valido.

Le fixture generano i file di esempio al volo nella ``tmp_path`` di pytest:
nessun binario di sistema (pandoc, tesseract) è richiesto per far girare i test.
"""

from __future__ import annotations

from email.message import EmailMessage
from pathlib import Path

import pytest

from extractors._base import ExtractionResult
from extractors._registry import (
    get_extractor,
    get_ocr_extractor,
    supported_extensions,
    supported_mimes,
)
from extractors.docx import DocxExtractor
from extractors.eml import EmlExtractor
from extractors.pdf import PdfExtractor
from extractors.pdf_ocr import PdfOcrExtractor
from extractors.plain import PlainExtractor
from extractors.xlsx import XlsxExtractor

# ---------------------------------------------------------------------------
# Fixture: generatori di file di esempio
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_txt(tmp_path: Path) -> Path:
    p = tmp_path / "appunti.txt"
    p.write_text("Riga uno.\nRiga due con accenti àèì.\n", encoding="utf-8")
    return p


@pytest.fixture
def sample_csv(tmp_path: Path) -> Path:
    p = tmp_path / "clienti.csv"
    p.write_text(
        "nome,settore,fatturato\nRossi Srl,edilizia,1200000\nBianchi Spa,manifattura,4500000\n",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def sample_md(tmp_path: Path) -> Path:
    p = tmp_path / "note.md"
    p.write_text("# Titolo\n\n- punto 1\n- punto 2\n", encoding="utf-8")
    return p


@pytest.fixture
def sample_pdf_minimal(tmp_path: Path) -> Path:
    """Fallback minimo: PDF "valido" cucito a mano se reportlab manca."""
    # PDF di una pagina con testo "Hello" — sufficiente per pdfplumber/pypdf.
    p = tmp_path / "mini.pdf"
    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 144]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\n"
        b"BT /F1 18 Tf 20 100 Td (Hello PDF) Tj ET\n"
        b"endstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000010 00000 n \n"
        b"0000000053 00000 n \n"
        b"0000000098 00000 n \n"
        b"0000000186 00000 n \n"
        b"0000000275 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\n"
        b"startxref\n330\n%%EOF\n"
    )
    p.write_bytes(pdf_bytes)
    return p


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    pytest.importorskip("docx")
    from docx import Document  # type: ignore[import-not-found]

    p = tmp_path / "relazione.docx"
    doc = Document()
    doc.add_heading("Relazione trimestrale", level=1)
    doc.add_paragraph("Paragrafo introduttivo con qualche dettaglio.")
    doc.add_heading("Sezione 2", level=2)
    doc.add_paragraph("Altro paragrafo.")
    # Tabella 2x2
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Voce"
    table.rows[0].cells[1].text = "Valore"
    table.rows[1].cells[0].text = "Fatturato"
    table.rows[1].cells[1].text = "100k"
    doc.save(str(p))
    return p


@pytest.fixture
def sample_xlsx(tmp_path: Path) -> Path:
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook  # type: ignore[import-not-found]

    p = tmp_path / "anagrafica.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Clienti"
    # Annotazioni libere prima della tabella
    ws["A1"] = "Nota: aggiornato 2024"
    # Header alla riga 3
    ws["A3"] = "Cliente"
    ws["B3"] = "Settore"
    ws["C3"] = "Fatturato"
    ws["A4"] = "Rossi"
    ws["B4"] = "edilizia"
    ws["C4"] = 100000
    ws["A5"] = "Bianchi"
    ws["B5"] = "industria"
    ws["C5"] = 200000
    # Foglio vuoto
    wb.create_sheet("Vuoto")
    wb.save(str(p))
    return p


@pytest.fixture
def sample_eml(tmp_path: Path) -> Path:
    p = tmp_path / "mail.eml"
    msg = EmailMessage()
    msg["From"] = "Mario Rossi <m.rossi@bianchi.it>"
    msg["To"] = "a.ferrari@bianchi.it"
    msg["Cc"] = "info@bianchi.it"
    msg["Subject"] = "Offerta Verdi Costruzioni"
    msg["Date"] = "Wed, 12 Mar 2024 10:23:00 +0100"
    msg["Message-ID"] = "<abc@bianchi.it>"
    msg.set_content("Ciao Anna,\n\nti allego l'offerta richiesta.\n\nSaluti,\nMario")
    msg.add_attachment(
        b"finto pdf content",
        maintype="application",
        subtype="pdf",
        filename="offerta.pdf",
    )
    p.write_bytes(bytes(msg))
    return p


# ---------------------------------------------------------------------------
# Plain
# ---------------------------------------------------------------------------


def test_plain_txt(sample_txt: Path) -> None:
    res = PlainExtractor().extract(sample_txt)
    assert isinstance(res, ExtractionResult)
    assert res.markdown.strip()
    assert "Riga uno" in res.markdown
    assert res.quality > 0.5
    assert res.metadata["format"] == "text"


def test_plain_md(sample_md: Path) -> None:
    res = PlainExtractor().extract(sample_md)
    assert res.metadata["format"] == "markdown"
    assert "# Titolo" in res.markdown


def test_plain_csv_markdown_table(sample_csv: Path) -> None:
    res = PlainExtractor().extract(sample_csv)
    assert res.metadata["format"] == "csv"
    lines = res.markdown.splitlines()
    # Header + separatore + 2 righe dati
    assert lines[0].startswith("| nome ")
    assert lines[1] == "| --- | --- | --- |"
    assert any("Rossi Srl" in line for line in lines[2:])
    assert any("Bianchi Spa" in line for line in lines[2:])


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def test_pdf_text_extraction(tmp_path: Path) -> None:
    """Usa reportlab se disponibile, altrimenti il PDF minimo cucito a mano."""
    try:
        from reportlab.pdfgen import canvas  # type: ignore[import-not-found]
    except ImportError:
        pytest.skip("reportlab non installato: skip PDF reale")

    p = tmp_path / "doc.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(72, 750, "Documento di prova - pagina 1")
    c.showPage()
    c.drawString(72, 750, "Documento di prova - pagina 2")
    c.showPage()
    c.save()

    res = PdfExtractor().extract(p)
    assert isinstance(res, ExtractionResult)
    assert "## Pagina 1" in res.markdown
    assert "## Pagina 2" in res.markdown
    assert res.metadata["pages"] == 2
    assert res.metadata["engine"] in {"pdfplumber", "pypdf"}
    assert res.quality > 0.0


def test_pdf_minimal_fallback(sample_pdf_minimal: Path) -> None:
    """Anche un PDF minimo non deve far crashare l'extractor."""
    res = PdfExtractor().extract(sample_pdf_minimal)
    assert isinstance(res, ExtractionResult)
    # Markdown sempre presente (anche se solo placeholder).
    assert res.markdown
    assert res.metadata["pages"] >= 1


# ---------------------------------------------------------------------------
# PDF OCR (graceful absence)
# ---------------------------------------------------------------------------


def test_pdf_ocr_graceful_without_tesseract(sample_pdf_minimal: Path) -> None:
    """Quando Tesseract non è installato, l'extractor non crasha e marca quality=0."""
    res = PdfOcrExtractor().extract(sample_pdf_minimal)
    assert isinstance(res, ExtractionResult)
    # In CI/sandbox tesseract non c'è → quality 0 e warning chiaro.
    # Se invece tesseract è installato accettiamo qualunque risultato non crash.
    import shutil

    if shutil.which("tesseract") is None:
        assert res.quality == 0.0
        assert any("Tesseract" in w for w in res.warnings)
        assert res.metadata.get("ocr_skipped") is True


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_docx_extraction(sample_docx: Path) -> None:
    res = DocxExtractor().extract(sample_docx)
    assert isinstance(res, ExtractionResult)
    assert res.markdown.strip()
    assert "Relazione trimestrale" in res.markdown
    assert "Fatturato" in res.markdown
    assert res.metadata["engine"] in {"pandoc", "python-docx"}


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def test_xlsx_extraction(sample_xlsx: Path) -> None:
    res = XlsxExtractor().extract(sample_xlsx)
    assert isinstance(res, ExtractionResult)
    # Header riconosciuto
    assert "## Clienti" in res.markdown
    assert "| Cliente | Settore | Fatturato |" in res.markdown
    assert "Rossi" in res.markdown
    assert "Bianchi" in res.markdown
    # Annotazione libera (riga 1) recuperata
    assert "aggiornato 2024" in res.markdown
    # Foglio vuoto skippato
    assert "## Vuoto" not in res.markdown
    assert res.metadata["sheets_count"] == 1


# ---------------------------------------------------------------------------
# EML
# ---------------------------------------------------------------------------


def test_eml_extraction(sample_eml: Path) -> None:
    res = EmlExtractor().extract(sample_eml)
    assert isinstance(res, ExtractionResult)
    assert res.markdown.startswith("---")
    assert "from:" in res.markdown
    assert "Offerta Verdi Costruzioni" in res.markdown
    assert "## Allegati" in res.markdown
    assert "offerta.pdf" in res.markdown
    assert "## Corpo" in res.markdown
    assert "ti allego" in res.markdown
    assert res.metadata["attachments"][0]["filename"] == "offerta.pdf"


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------


def test_registry_by_mime() -> None:
    assert isinstance(get_extractor("application/pdf", ".pdf"), PdfExtractor)
    assert isinstance(get_extractor("message/rfc822", ".eml"), EmlExtractor)
    assert isinstance(
        get_extractor(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".docx",
        ),
        DocxExtractor,
    )
    assert isinstance(
        get_extractor(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xlsx",
        ),
        XlsxExtractor,
    )
    assert isinstance(get_extractor("text/plain", ".txt"), PlainExtractor)
    assert isinstance(get_extractor("text/csv", ".csv"), PlainExtractor)


def test_registry_fallback_to_extension() -> None:
    # Mime sconosciuto ma estensione nota → trovato tramite estensione.
    assert isinstance(get_extractor(None, ".pdf"), PdfExtractor)
    assert isinstance(get_extractor("application/octet-stream", ".docx"), DocxExtractor)


def test_registry_unknown_returns_none() -> None:
    assert get_extractor(None, ".dwg") is None
    assert get_extractor("application/x-foo-bar", "") is None


def test_registry_ocr_not_in_primary_lookup() -> None:
    """``pdf_ocr`` non deve essere selezionato dal lookup automatico per ``application/pdf``."""
    primary = get_extractor("application/pdf", ".pdf")
    assert not isinstance(primary, PdfOcrExtractor)
    # Ma deve essere disponibile esplicitamente.
    assert isinstance(get_ocr_extractor(), PdfOcrExtractor)


def test_registry_supported_lists_nonempty() -> None:
    assert "application/pdf" in supported_mimes()
    assert ".docx" in supported_extensions()
