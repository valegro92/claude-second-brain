"""Genera un dataset pilota realistico per i test end-to-end.

Output: una cartella ``_pilot/inbox/`` con 40-60 file misti (PDF, DOCX, XLSX,
EML, TXT) riconducibili a 3 clienti finti e 2 fornitori, distribuiti su
diverse età, con qualche "trappola" tipica (versioni multiple, copie, cartelle
``_OLD_NON_USARE/``) e qualche binario da escludere.

Uso CLI::

    python -m tests.fixtures.build_pilot_dataset --output _pilot/inbox/

Il comando è idempotente: se la cartella esiste viene cancellata e rigenerata
da capo. Il risultato deve essere riproducibile (seed fisso) così i test E2E
possono asserire numeri stabili.

Dipendenze:
  * ``python-docx`` (sempre presente, parte del toolkit)
  * ``openpyxl`` (sempre presente)
  * ``reportlab`` opzionale: se installata genera PDF "veri" con testo; se
    manca, ricade su un PDF minimo cucito a mano (come in
    ``tests/test_extractors.py``).
"""

from __future__ import annotations

import argparse
import logging
import os
import random
import shutil
from collections.abc import Callable
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from email.message import EmailMessage
from pathlib import Path

logger = logging.getLogger(__name__)


# --- costanti -------------------------------------------------------------

# Seed fisso: i test E2E asseriscono conteggi.
RANDOM_SEED = 42

# Numero target di file generati (range "40-60" dal brief): scegliamo 48 stabile.
TARGET_FILE_COUNT = 48

# Clienti finti.
CLIENTI = ["Rossi Srl", "Verdi Costruzioni", "Bianchi SpA"]

# Fornitori finti.
FORNITORI = ["Acciai Lombardi", "Logistica Veneta"]

# Mtime "epoche": il categorizer rules usa l'età, qui generiamo un mix.
EPOCHE: dict[str, datetime] = {
    "recente_2024": datetime(2024, 6, 15, tzinfo=UTC),
    "medio_2022": datetime(2022, 3, 10, tzinfo=UTC),
    "vecchio_2020": datetime(2020, 9, 1, tzinfo=UTC),
    "arcaico_2018": datetime(2018, 2, 14, tzinfo=UTC),
}

# Mini-PDF cucito a mano (vedi tests/test_extractors.py).
_MINIMAL_PDF_BYTES = (
    b"%PDF-1.4\n"
    b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
    b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 144]"
    b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
    b"4 0 obj<</Length 44>>stream\n"
    b"BT /F1 18 Tf 20 100 Td (Hello PDF) Tj ET\n"
    b"endstream endobj\n"
    b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
    b"xref\n0 6\n"
    b"0000000000 65535 f \n"
    b"0000000010 00000 n \n"
    b"0000000053 00000 n \n"
    b"0000000098 00000 n \n"
    b"0000000186 00000 n \n"
    b"0000000275 00000 n \n"
    b"trailer<</Size 6/Root 1 0 R>>\n"
    b"startxref\n330\n%%EOF\n"
)


# --- modelli -------------------------------------------------------------


@dataclass
class GeneratedFile:
    """Output di un generatore: usato per il logging finale."""

    path: Path
    soggetto: str  # "cliente:Rossi Srl", "fornitore:Acciai Lombardi", ...
    epoca: str  # chiave EPOCHE
    tipo: str  # pdf | docx | xlsx | eml | txt | bin


# --- helper generici -----------------------------------------------------


def _slug(name: str) -> str:
    return name.lower().replace(" ", "-").replace(".", "").strip("-")


def _set_mtime(path: Path, dt: datetime) -> None:
    """Forza mtime/atime sul file."""
    ts = dt.timestamp()
    os.utime(path, (ts, ts))


def _write_and_age(path: Path, content: bytes | str, epoca_dt: datetime) -> Path:
    """Scrive un file e applica un mtime "invecchiato"."""
    path.parent.mkdir(parents=True, exist_ok=True)
    if isinstance(content, str):
        path.write_text(content, encoding="utf-8")
    else:
        path.write_bytes(content)
    _set_mtime(path, epoca_dt)
    return path


# --- generatori per tipo --------------------------------------------------


def _generate_pdf(path: Path, titolo: str, body: str, epoca_dt: datetime) -> Path:
    """Genera un PDF: usa reportlab se disponibile, altrimenti minimal-PDF.

    Senza reportlab i PDF condividerebbero lo stesso payload byte-per-byte
    (collisione di sha256, problemi per i test E2E che asseriscono N
    estrazioni distinte). Aggiungiamo quindi un commento PDF univoco
    (riga ``%`` ignorata dai parser) col path + titolo, in coda al file.
    Resta un PDF valido, ma con sha distinto.
    """
    try:
        from reportlab.pdfgen import canvas  # type: ignore[import-not-found]
    except ImportError:
        path.parent.mkdir(parents=True, exist_ok=True)
        salt = f"%uniq:{path.name}:{titolo}:{epoca_dt.isoformat()}\n".encode()
        # Il commento `%` PDF va prima di `%%EOF` per non rompere il trailer.
        eof = b"%%EOF\n"
        payload = _MINIMAL_PDF_BYTES
        if payload.endswith(eof):
            payload = payload[: -len(eof)] + salt + eof
        else:
            payload = payload + b"\n" + salt
        return _write_and_age(path, payload, epoca_dt)

    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path))
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 760, titolo[:80])
    c.setFont("Helvetica", 10)
    y = 730
    for line in body.splitlines():
        if y < 60:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = 760
        c.drawString(72, y, line[:100])
        y -= 14
    c.save()
    _set_mtime(path, epoca_dt)
    return path


def _generate_docx(path: Path, titolo: str, body: str, epoca_dt: datetime) -> Path:
    """Genera un DOCX con `python-docx` (sempre presente)."""
    from docx import Document  # type: ignore[import-not-found]

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(titolo[:80], level=1)
    for paragrafo in body.split("\n\n"):
        doc.add_paragraph(paragrafo.strip())
    doc.save(str(path))
    _set_mtime(path, epoca_dt)
    return path


def _generate_xlsx(
    path: Path,
    sheet_name: str,
    header: list[str],
    rows: list[list[str | int]],
    epoca_dt: datetime,
) -> Path:
    """Genera un XLSX con `openpyxl`."""
    from openpyxl import Workbook  # type: ignore[import-not-found]

    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31]
    for c, h in enumerate(header, start=1):
        ws.cell(row=1, column=c, value=h)
    for r, row in enumerate(rows, start=2):
        for c, val in enumerate(row, start=1):
            ws.cell(row=r, column=c, value=val)
    wb.save(str(path))
    _set_mtime(path, epoca_dt)
    return path


def _generate_eml(
    path: Path,
    mittente: str,
    destinatario: str,
    oggetto: str,
    corpo: str,
    epoca_dt: datetime,
    allegato_name: str | None = None,
    allegato_bytes: bytes | None = None,
) -> Path:
    """Genera un EML, opzionalmente con un allegato."""
    path.parent.mkdir(parents=True, exist_ok=True)
    msg = EmailMessage()
    msg["From"] = mittente
    msg["To"] = destinatario
    msg["Subject"] = oggetto
    msg["Date"] = epoca_dt.strftime("%a, %d %b %Y %H:%M:%S +0000")
    msg["Message-ID"] = f"<pilot-{abs(hash((oggetto, epoca_dt.isoformat()))):x}@pilot.test>"
    msg.set_content(corpo)
    if allegato_name and allegato_bytes:
        # Determina maintype/subtype dall'estensione.
        ext = Path(allegato_name).suffix.lower()
        maintype, subtype = {
            ".pdf": ("application", "pdf"),
            ".docx": (
                "application",
                "vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            ".xlsx": (
                "application",
                "vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ),
            ".txt": ("text", "plain"),
        }.get(ext, ("application", "octet-stream"))
        msg.add_attachment(
            allegato_bytes, maintype=maintype, subtype=subtype, filename=allegato_name
        )
    path.write_bytes(bytes(msg))
    _set_mtime(path, epoca_dt)
    return path


def _generate_txt(path: Path, body: str, epoca_dt: datetime) -> Path:
    return _write_and_age(path, body, epoca_dt)


# --- contenuti finti ------------------------------------------------------


def _testo_offerta(cliente: str, anno: int) -> tuple[str, str]:
    titolo = f"Offerta commerciale {cliente} - {anno}"
    body = (
        f"Spett.le {cliente},\n\n"
        f"in riferimento alla Vostra richiesta del mese scorso, siamo lieti di "
        f"sottoporre la nostra proposta commerciale per la fornitura della "
        f"commessa {anno}/00{random.randint(1, 99)}.\n\n"
        f"Le condizioni economiche sono dettagliate nella tabella in calce.\n"
        f"Validità offerta: 30 giorni.\n\n"
        f"Cordiali saluti,\n"
        f"Ufficio Commerciale"
    )
    return titolo, body


def _testo_contratto(controparte: str, anno: int) -> tuple[str, str]:
    titolo = f"Contratto di fornitura {controparte} - {anno}"
    body = (
        f"CONTRATTO DI FORNITURA\n\n"
        f"Tra le sotto-elencate parti, in data {anno}-{random.randint(1, 12):02d}-15:\n\n"
        f"- Committente: La Nostra Azienda Srl\n"
        f"- Fornitore: {controparte}\n\n"
        f"Oggetto: fornitura di servizi e prodotti come da capitolato allegato. "
        f"Importo annuo: euro {random.randint(10, 500)}.000,00.\n\n"
        f"Le parti si impegnano a osservare gli accordi sotto-elencati.\n"
        f"Firmato in originale dalle parti."
    )
    return titolo, body


def _testo_ordine(cliente: str, anno: int) -> tuple[str, str]:
    titolo = f"Ordine {cliente}"
    body = (
        f"ORDINE n. {anno}/{random.randint(1, 999):03d}\n\n"
        f"Cliente: {cliente}\n"
        f"Data: {anno}-{random.randint(1, 12):02d}-{random.randint(1, 28):02d}\n\n"
        f"Riferimento offerta: vedi documento allegato.\n"
        f"Termini di pagamento: 30 giorni data fattura.\n"
    )
    return titolo, body


def _anagrafica_xlsx_rows(soggetti: list[str]) -> list[list[str | int]]:
    return [
        [
            s,
            random.choice(["edilizia", "manifattura", "servizi", "logistica"]),
            random.randint(500, 5000) * 1000,
        ]
        for s in soggetti
    ]


# --- generatore principale ------------------------------------------------


def _setup_random() -> None:
    random.seed(RANDOM_SEED)


def _ensure_clean_output(output_dir: Path) -> None:
    if output_dir.exists():
        logger.info("Cartella esistente: cancello e ricreo (%s)", output_dir)
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)


def _add(generated: list[GeneratedFile], path: Path, soggetto: str, epoca: str, tipo: str) -> None:
    generated.append(GeneratedFile(path=path, soggetto=soggetto, epoca=epoca, tipo=tipo))


def build_dataset(output_dir: Path) -> list[GeneratedFile]:
    """Genera il dataset completo in ``output_dir``. Idempotente.

    Ritorna la lista dei file generati (per logging / debug / asserzioni nei test).
    """
    _setup_random()
    _ensure_clean_output(output_dir)
    generated: list[GeneratedFile] = []

    clienti_dir = output_dir / "clienti"
    fornitori_dir = output_dir / "fornitori"
    commerciale_dir = output_dir / "commerciale"
    archivio_old_dir = output_dir / "_OLD_NON_USARE"

    # --- 1. Per ogni cliente: 1 cartella, 1 PDF offerta recente, 1 DOCX ordine,
    #        1 XLSX anagrafica, qualche EML.
    for cliente in CLIENTI:
        slug = _slug(cliente)
        cdir = clienti_dir / slug
        # Offerta 2024 (recente)
        titolo, body = _testo_offerta(cliente, 2024)
        p = _generate_pdf(cdir / f"{slug}_offerta_2024.pdf", titolo, body, EPOCHE["recente_2024"])
        _add(generated, p, f"cliente:{cliente}", "recente_2024", "pdf")

        # Ordine 2022 (medio)
        titolo, body = _testo_ordine(cliente, 2022)
        p = _generate_docx(cdir / f"{slug}_ordine_2022.docx", titolo, body, EPOCHE["medio_2022"])
        _add(generated, p, f"cliente:{cliente}", "medio_2022", "docx")

        # Contratto 2020 (vecchio)
        titolo, body = _testo_contratto(cliente, 2020)
        p = _generate_pdf(cdir / f"{slug}_contratto_2020.pdf", titolo, body, EPOCHE["vecchio_2020"])
        _add(generated, p, f"cliente:{cliente}", "vecchio_2020", "pdf")

        # Anagrafica xlsx (medio)
        p = _generate_xlsx(
            cdir / f"{slug}_referenti.xlsx",
            sheet_name="Referenti",
            header=["Nome", "Ruolo", "Email"],
            rows=[
                ["Mario Bianchi", "Acquisti", f"acquisti@{slug}.it"],
                ["Anna Verdi", "Amministrazione", f"amm@{slug}.it"],
            ],
            epoca_dt=EPOCHE["medio_2022"],
        )
        _add(generated, p, f"cliente:{cliente}", "medio_2022", "xlsx")

        # 2 email scambiate col cliente (una recente, una vecchia, una con allegato)
        eml = _generate_eml(
            cdir / f"email_{slug}_001.eml",
            mittente=f"acquisti@{slug}.it",
            destinatario="commerciale@nostraazienda.it",
            oggetto=f"Richiesta offerta {cliente}",
            corpo=(
                f"Buongiorno,\n\nVi chiediamo gentilmente di formulare una "
                f"nuova offerta per la fornitura annuale. Trovate in allegato "
                f"il capitolato tecnico.\n\nGrazie e cordiali saluti,\n"
                f"Ufficio Acquisti {cliente}"
            ),
            epoca_dt=EPOCHE["recente_2024"],
            allegato_name=f"capitolato_{slug}.pdf",
            allegato_bytes=_MINIMAL_PDF_BYTES,
        )
        _add(generated, eml, f"cliente:{cliente}", "recente_2024", "eml")

        eml2 = _generate_eml(
            cdir / f"email_{slug}_002.eml",
            mittente="commerciale@nostraazienda.it",
            destinatario=f"acquisti@{slug}.it",
            oggetto=f"Riepilogo riunione del 2022 - {cliente}",
            corpo=(
                "Gentile cliente,\n\nFacciamo seguito alla riunione "
                "del 2022 per riepilogare i punti concordati: tempi di "
                "consegna, listino prezzi, modalità di pagamento.\n\n"
                "Restiamo a disposizione."
            ),
            epoca_dt=EPOCHE["medio_2022"],
        )
        _add(generated, eml2, f"cliente:{cliente}", "medio_2022", "eml")

    # --- 2. Per ogni fornitore: 1 cartella, 1 contratto + 1 fattura mock + 1 email
    for fornitore in FORNITORI:
        slug = _slug(fornitore)
        fdir = fornitori_dir / slug

        # Contratto 2020 (vecchio)
        titolo, body = _testo_contratto(fornitore, 2020)
        p = _generate_pdf(fdir / f"{slug}_contratto.pdf", titolo, body, EPOCHE["vecchio_2020"])
        _add(generated, p, f"fornitore:{fornitore}", "vecchio_2020", "pdf")

        # "Fattura" come TXT (semplificazione)
        p = _generate_txt(
            fdir / f"{slug}_fattura_2024_001.txt",
            body=(
                f"FATTURA {fornitore}\n"
                f"Numero: 2024/{random.randint(100, 999)}\n"
                f"Importo: euro {random.randint(100, 5000)},00\n"
                f"Causale: fornitura ricorrente.\n"
            ),
            epoca_dt=EPOCHE["recente_2024"],
        )
        _add(generated, p, f"fornitore:{fornitore}", "recente_2024", "txt")

        # Email di sollecito
        eml = _generate_eml(
            fdir / f"email_{slug}_sollecito.eml",
            mittente=f"amministrazione@{slug}.it",
            destinatario="amm@nostraazienda.it",
            oggetto=f"Sollecito pagamento - {fornitore}",
            corpo=(
                f"Spettabile cliente,\n\nVi sollecitiamo il saldo della "
                f"fattura {random.randint(100, 999)} emessa nei mesi scorsi.\n\n"
                f"Cordiali saluti,\n{fornitore}"
            ),
            epoca_dt=EPOCHE["medio_2022"],
        )
        _add(generated, eml, f"fornitore:{fornitore}", "medio_2022", "eml")

    # --- 3. Cartella commerciale: anagrafica generale e qualche file libero
    p = _generate_xlsx(
        commerciale_dir / "anagrafica_clienti.xlsx",
        sheet_name="Clienti",
        header=["Ragione sociale", "Settore", "Fatturato"],
        rows=_anagrafica_xlsx_rows(CLIENTI),
        epoca_dt=EPOCHE["recente_2024"],
    )
    _add(generated, p, "anagrafica", "recente_2024", "xlsx")

    p = _generate_txt(
        commerciale_dir / "appunti_riunione_settimanale.txt",
        body="Punti discussi: offerte aperte, scadenze pagamento, contatti caldi.\n",
        epoca_dt=EPOCHE["recente_2024"],
    )
    _add(generated, p, "commerciale-generico", "recente_2024", "txt")

    p = _generate_docx(
        commerciale_dir / "procedura_offerta.docx",
        titolo="Procedura per la preparazione di un'offerta",
        body=(
            "1. Ricezione della richiesta dal cliente.\n\n"
            "2. Verifica disponibilità tecnica e tempi.\n\n"
            "3. Preparazione offerta su template ufficiale.\n\n"
            "4. Validazione dal responsabile commerciale.\n\n"
            "5. Invio formale via PEC al cliente."
        ),
        epoca_dt=EPOCHE["medio_2022"],
    )
    _add(generated, p, "commerciale-generico", "medio_2022", "docx")

    # --- 4. TRAPPOLE: 3 file _v1/_v2/_v3/_FINAL nel cliente Rossi
    slug_rossi = _slug("Rossi Srl")
    trap_dir = clienti_dir / slug_rossi / "offerte_in_lavorazione"
    for n, suffix in enumerate(("_v1", "_v2", "_v3"), start=1):
        titolo, body = _testo_offerta("Rossi Srl", 2024)
        p = _generate_docx(
            trap_dir / f"offerta_speciale{suffix}.docx",
            titolo=f"{titolo} ({suffix})",
            body=body,
            epoca_dt=EPOCHE["recente_2024"] - timedelta(days=10 - n),
        )
        _add(generated, p, "cliente:Rossi Srl", "recente_2024", "docx")
    # Versione "FINAL"
    titolo, body = _testo_offerta("Rossi Srl", 2024)
    p = _generate_docx(
        trap_dir / "offerta_speciale_FINAL.docx",
        titolo=f"{titolo} (FINAL)",
        body=body,
        epoca_dt=EPOCHE["recente_2024"],
    )
    _add(generated, p, "cliente:Rossi Srl", "recente_2024", "docx")

    # --- 5. TRAPPOLE: 2 file "copia di" nel cliente Verdi
    slug_verdi = _slug("Verdi Costruzioni")
    vdir = clienti_dir / slug_verdi
    p = _generate_pdf(
        vdir / "copia di preventivo.pdf",
        titolo="Copia di preventivo",
        body="Vedi documento originale.",
        epoca_dt=EPOCHE["medio_2022"],
    )
    _add(generated, p, "cliente:Verdi Costruzioni", "medio_2022", "pdf")
    p = _generate_pdf(
        vdir / "copia di copia di preventivo.pdf",
        titolo="Copia di copia di preventivo",
        body="Duplicato del duplicato.",
        epoca_dt=EPOCHE["medio_2022"],
    )
    _add(generated, p, "cliente:Verdi Costruzioni", "medio_2022", "pdf")

    # --- 6. Cartella _OLD_NON_USARE/: roba archivistica
    p = _generate_txt(
        archivio_old_dir / "vecchi_listini_2018.txt",
        body="Listino prezzi 2018. NON usare per nuove offerte.\n",
        epoca_dt=EPOCHE["arcaico_2018"],
    )
    _add(generated, p, "archivio", "arcaico_2018", "txt")
    p = _generate_docx(
        archivio_old_dir / "manuale_procedure_v0.docx",
        titolo="Manuale procedure (obsoleto)",
        body="Bozza iniziale del manuale, sostituita dalla v2.\n\nNon più valida.",
        epoca_dt=EPOCHE["arcaico_2018"],
    )
    _add(generated, p, "archivio", "arcaico_2018", "docx")
    p = _generate_pdf(
        archivio_old_dir / "vecchio_contratto_2018.pdf",
        titolo="Contratto 2018 (scaduto)",
        body="Contratto scaduto e sostituito.",
        epoca_dt=EPOCHE["arcaico_2018"],
    )
    _add(generated, p, "archivio", "arcaico_2018", "pdf")

    # --- 7. FILE DA ESCLUDERE: 1 .dwg "binario" (stub) + 1 .zip vuoto
    #     (i filtri_globali nel config del cliente li scartano).
    p = output_dir / "disegno_tecnico.dwg"
    p.write_bytes(b"\x00\x01\x02FAKE-DWG-HEADER" + b"\x00" * 1000)
    _set_mtime(p, EPOCHE["vecchio_2020"])
    _add(generated, p, "binario-escluso", "vecchio_2020", "bin")

    p = output_dir / "backup_grande.zip"
    p.write_bytes(b"PK\x03\x04")  # ZIP minimo, simulazione "binario da escludere"
    _set_mtime(p, EPOCHE["vecchio_2020"])
    _add(generated, p, "binario-escluso", "vecchio_2020", "bin")

    # --- 8. Riempitivi per arrivare al target.
    riempitivi_dir = output_dir / "varie"
    while len(generated) < TARGET_FILE_COUNT:
        idx = len(generated) + 1
        soggetto = random.choice(CLIENTI + FORNITORI)
        epoca_key = random.choice(list(EPOCHE.keys()))
        scelta: Callable[[], tuple[Path, str]] = random.choice(
            [
                lambda: (
                    riempitivi_dir / f"nota_{idx:03d}.txt",
                    "txt",
                ),
                lambda: (
                    riempitivi_dir / f"memo_{idx:03d}.docx",
                    "docx",
                ),
            ]
        )
        target_path, tipo = scelta()
        if tipo == "txt":
            p = _generate_txt(
                target_path,
                body=f"Nota {idx} riferita a {soggetto}. Epoca: {epoca_key}.\n",
                epoca_dt=EPOCHE[epoca_key],
            )
        else:
            p = _generate_docx(
                target_path,
                titolo=f"Memo {idx} - {soggetto}",
                body=f"Memo interno riferito a {soggetto}.\n\nEpoca: {epoca_key}.",
                epoca_dt=EPOCHE[epoca_key],
            )
        _add(generated, p, f"riempitivo:{soggetto}", epoca_key, tipo)

    logger.info(
        "Dataset pilota generato in %s: %d file (target %d).",
        output_dir,
        len(generated),
        TARGET_FILE_COUNT,
    )
    return generated


# --- CLI -----------------------------------------------------------------


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Genera il dataset pilota per i test E2E.")
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("_pilot/inbox"),
        help="Cartella di output (default: _pilot/inbox/). Ricreata da zero.",
    )
    parser.add_argument("-v", "--verbose", action="store_true", help="Logging dettagliato")
    args = parser.parse_args(argv)
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)s %(name)s: %(message)s",
    )

    files = build_dataset(args.output)
    print(f"Generati {len(files)} file in {args.output}")
    # Riepilogo per tipo
    by_tipo: dict[str, int] = {}
    for f in files:
        by_tipo[f.tipo] = by_tipo.get(f.tipo, 0) + 1
    for tipo, n in sorted(by_tipo.items()):
        print(f"  {tipo}: {n}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
