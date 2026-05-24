"""
Parser DOCX → testo via python-docx.

Concatena, in ordine:
1. Header di ogni sezione (prefisso ``## Header:``), una sola volta per header
   distinto, perché python-docx ritorna lo stesso oggetto su sezioni con header
   ereditato.
2. Paragrafi del body + celle di tabella (separati da newline, celle da
   ``" | "``).
3. Footer di ogni sezione (prefisso ``## Footer:``), come gli header.

Header/footer contengono spesso ragione sociale, P.IVA, riferimenti contratto
sulle offerte italiane: vale la pena estrarli per migliorare la qualità delle
entity extraction downstream.

Anche i Google Docs nativi finiscono qui: il connettore Drive li esporta come
DOCX prima di passarli a questo parser.
"""

from __future__ import annotations

import io
from pathlib import Path

import docx  # python-docx
from docx.opc.exceptions import PackageNotFoundError

from custodia_cli.connectors.base import ParserError


def _collect_section_block_texts(paragraphs_iter) -> list[str]:  # type: ignore[no-untyped-def]
    """Concatena testo non-vuoto dai paragrafi di header/footer."""
    out: list[str] = []
    for p in paragraphs_iter:
        text = p.text.strip()
        if text:
            out.append(text)
    return out


def parse_docx(content: bytes | Path) -> str:
    """Estrae header + paragrafi + celle tabella + footer da un file DOCX.

    Args:
        content: bytes del file o ``Path`` a un DOCX su disco.

    Returns:
        Testo concatenato (header/body/footer separati da ``\\n``). Header e
        footer sono prefissati da ``## Header:`` / ``## Footer:`` per
        permettere agli LLM downstream di distinguerli dal body.

    Raises:
        ParserError: se il file non è un DOCX valido / package corrotto.
    """
    # Difesa zip-bomb prima di passare il file a python-docx.
    # Import locale per evitare ciclo (parsers.__init__ importa questo modulo).
    from custodia_cli.connectors.parsers import check_zip_uncompressed_size

    check_zip_uncompressed_size(content)

    stream: io.BytesIO | Path
    if isinstance(content, Path):
        stream = content
    else:
        stream = io.BytesIO(content)

    try:
        document = docx.Document(stream)
    except PackageNotFoundError as exc:
        raise ParserError(f"DOCX non valido o corrotto: {exc}") from exc
    except Exception as exc:  # noqa: BLE001
        raise ParserError(f"Errore inatteso aprendo DOCX: {exc}") from exc

    parts: list[str] = []

    # 1. Headers (deduplicati per testo: sezioni con header ereditato
    #    espongono lo stesso testo, evitiamo ripetizioni).
    seen_headers: set[str] = set()
    header_lines: list[str] = []
    for section in document.sections:
        try:
            lines = _collect_section_block_texts(section.header.paragraphs)
        except Exception:  # noqa: BLE001 — header malformato non blocca il body
            lines = []
        for line in lines:
            if line not in seen_headers:
                seen_headers.add(line)
                header_lines.append(line)
    if header_lines:
        parts.append("## Header:")
        parts.extend(header_lines)

    # 2. Body: paragrafi.
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # 2b. Body: tabelle.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    # 3. Footers (stessa logica di dedup).
    seen_footers: set[str] = set()
    footer_lines: list[str] = []
    for section in document.sections:
        try:
            lines = _collect_section_block_texts(section.footer.paragraphs)
        except Exception:  # noqa: BLE001
            lines = []
        for line in lines:
            if line not in seen_footers:
                seen_footers.add(line)
                footer_lines.append(line)
    if footer_lines:
        parts.append("## Footer:")
        parts.extend(footer_lines)

    return "\n".join(parts)


__all__ = ["parse_docx"]
