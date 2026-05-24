"""
Genera una cartella sintetica di N file (default 5000) con proporzioni
realistiche per la cartella di lavoro di un consulente PMI.

Composizione (5000 file totali, valore di default):
- 40% (2000): PDF brevi (1-3 pagine, testo plain in italiano)
- 25% (1250): DOCX (offerte, contratti, mail esportate)
- 15% (750):  XLSX (listini, anagrafica)
- 10% (500):  TXT/MD (note, README)
- 10% (500):  rumore da SKIPPARE (immagini finte, video finti, archivi, eseguibili)

Struttura cartelle (esempio con 5000 file):
  stress-corpus/
    clienti/
      acme-spa/
      bianchi-srl/
      ...
    fornitori/
    commesse/
    archivio/
      2020/
      2021/
      ...
    media/
      foto/         (skip: HEIC/JPG)
      video/        (skip: MOV/MP4)
    build/          (skip: node_modules + lock files)

Strategia di velocità: genera UNA template binaria di ogni tipo (PDF/DOCX/XLSX)
usando pypdf/python-docx/openpyxl, poi la copia (shutil.copy) per produrre N
file con path/source_id diversi. Il contenuto testuale è (semi-)uguale, ma per
il connector ciò che conta è il path (source_id = sha1(path)) e l'estrazione
del testo funziona correttamente.

Idempotenza: se la directory contiene già almeno ``--count`` file, l'esecuzione
è no-op. Per forzare la rigenerazione, cancellare la directory.

CLI standalone:
    python -m tests.fixtures.stress.generate --target /tmp/custodia-stress
    python tests/fixtures/stress/generate.py --count 5000
"""

from __future__ import annotations

import argparse
import io
import random
import shutil
import string
import sys
import tempfile
from pathlib import Path
from typing import Callable

# ---------------------------------------------------------------------------
# Template generators (1 PDF, 1 DOCX, 1 XLSX) — chiamati una sola volta
# ---------------------------------------------------------------------------


def _build_pdf_template() -> bytes:
    """PDF minimale valido (1 pagina) con pypdf. Ritorna i bytes."""
    from pypdf import PdfWriter
    from pypdf.generic import (
        DictionaryObject,
        NameObject,
        StreamObject,
    )

    writer = PdfWriter()
    page = writer.add_blank_page(width=595, height=842)
    text_lines = [
        "OFFERTA COMMERCIALE",
        "Cliente: Cliente Test SRL",
        "P.IVA: 00000000000",
        "Importo totale: EUR 4.250,00",
        "Descrizione: fornitura standard per consulenza",
    ]
    font_obj = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font_obj)  # type: ignore[attr-defined]
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream_lines = ["BT", "/F1 11 Tf"]
    y = 800
    for line in text_lines:
        safe = line.replace("(", r"\(").replace(")", r"\)")
        stream_lines.append(f"1 0 0 1 50 {y} Tm ({safe}) Tj")
        y -= 20
    stream_lines.append("ET")
    stream = StreamObject()
    stream.set_data("\n".join(stream_lines).encode("latin-1"))
    stream_ref = writer._add_object(stream)  # type: ignore[attr-defined]
    page[NameObject("/Contents")] = stream_ref

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _build_docx_template() -> bytes:
    """DOCX minimale con python-docx. Ritorna i bytes."""
    import docx

    doc = docx.Document()
    doc.add_heading("Offerta commerciale", level=1)
    doc.add_paragraph("Cliente: Bianchi Impianti SpA")
    doc.add_paragraph("Riferimento: OFF-STRESS-001")
    doc.add_paragraph(
        "Confermiamo disponibilità a fornire le quantità richieste."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx_template() -> bytes:
    """XLSX minimale con openpyxl. Ritorna i bytes."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Listino"
    ws.append(["Codice", "Descrizione", "Prezzo EUR"])
    ws.append(["LAM-12", "Lamiera zincata 1.2mm", 18.50])
    ws.append(["VL-DN50", "Valvola sfera DN50", 145.00])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Layout
# ---------------------------------------------------------------------------


_CLIENTI = [
    "acme-spa",
    "bianchi-srl",
    "rossetto-laminazioni",
    "torrelli-impianti",
    "verdi-costruzioni",
    "neri-componenti",
    "ferrari-meccanica",
    "marini-arredi",
    "decathlon",
    "ceteas",
]

_FORNITORI = [
    "fornitore-alpha",
    "fornitore-beta",
    "fornitore-gamma",
    "fornitore-delta",
    "fornitore-epsilon",
]

_COMMESSE = [
    "comm-2023-01",
    "comm-2023-02",
    "comm-2024-01",
    "comm-2024-02",
    "comm-2024-03",
    "comm-2025-01",
]

_ARCHIVIO_ANNI = ["2020", "2021", "2022", "2023"]


def _random_basename(rng: random.Random, prefix: str) -> str:
    """Nome file plausibile, unico abbastanza da non collidere su 5000 file."""
    suffix = "".join(rng.choices(string.ascii_lowercase + string.digits, k=8))
    return f"{prefix}-{suffix}"


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------


def _count_files(root: Path) -> int:
    """Conteggio file ricorsivo. Ritorna 0 se root non esiste."""
    if not root.exists():
        return 0
    return sum(1 for p in root.rglob("*") if p.is_file())


def generate_corpus(
    target_dir: Path,
    total_files: int = 5000,
    seed: int = 42,
    *,
    verbose: bool = False,
) -> Path:
    """Genera (o riusa) la cartella sintetica.

    Idempotente: se ``target_dir`` esiste e contiene ``>= total_files`` file,
    ritorna subito senza fare nulla. Altrimenti popola fino a raggiungere il
    target.

    Args:
        target_dir: cartella di destinazione (creata se mancante).
        total_files: numero di file da generare (default 5000).
        seed: seed RNG per layout deterministico.
        verbose: se True logga i progressi su stdout.

    Returns:
        Path della directory popolata.
    """
    target_dir = target_dir.expanduser().resolve()
    target_dir.mkdir(parents=True, exist_ok=True)

    existing = _count_files(target_dir)
    if existing >= total_files:
        if verbose:
            print(f"[generate] corpus già pronto: {existing} file in {target_dir}")
        return target_dir

    rng = random.Random(seed)

    # Prebuild template bytes (una volta, riusate per N file).
    if verbose:
        print("[generate] building template binaries (pypdf, python-docx, openpyxl)…")
    pdf_bytes = _build_pdf_template()
    docx_bytes = _build_docx_template()
    xlsx_bytes = _build_xlsx_template()

    # Quote per categoria (proporzioni realistiche).
    quotas: dict[str, int] = {
        "pdf": int(total_files * 0.40),     # 2000
        "docx": int(total_files * 0.25),    # 1250
        "xlsx": int(total_files * 0.15),    # 750
        "txt": int(total_files * 0.10),     # 500
        "noise": int(total_files * 0.10),   # 500
    }
    # Aggiusta per arrotondamento.
    diff = total_files - sum(quotas.values())
    quotas["pdf"] += diff

    # Genera path layout: per ogni file, scegli una sottocartella plausibile.
    def _pick_cliente_dir() -> Path:
        cli = rng.choice(_CLIENTI)
        return target_dir / "clienti" / cli

    def _pick_fornitore_dir() -> Path:
        return target_dir / "fornitori" / rng.choice(_FORNITORI)

    def _pick_commessa_dir() -> Path:
        return target_dir / "commesse" / rng.choice(_COMMESSE)

    def _pick_archivio_dir() -> Path:
        return target_dir / "archivio" / rng.choice(_ARCHIVIO_ANNI)

    # Distribuzione path per file di "lavoro" (PDF/DOCX/XLSX/TXT).
    def _work_dir() -> Path:
        roll = rng.random()
        if roll < 0.55:
            return _pick_cliente_dir()
        if roll < 0.75:
            return _pick_commessa_dir()
        if roll < 0.90:
            return _pick_fornitore_dir()
        return _pick_archivio_dir()

    def _write_binary(target: Path, content: bytes) -> None:
        target.parent.mkdir(parents=True, exist_ok=True)
        with target.open("wb") as fh:
            fh.write(content)

    def _write_text(target: Path, content: str) -> None:
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding="utf-8")

    n_created = 0

    # Per evitare che il manifest "rename detection" collassi tutti i file
    # con stesso template binario sullo stesso source_id (l'hash content è
    # uguale), iniettiamo un suffisso univoco di pochi byte alla fine del
    # contenuto. Per PDF/DOCX/XLSX (formati container) basta appendere un
    # commento ASCII dopo i bytes del container — i parser sono robusti agli
    # extra bytes terminali e l'hash del file resta unico.

    def _unique_suffix(idx: int, tag: str) -> bytes:
        return f"\n% custodia-stress-{tag}-{idx}\n".encode("ascii")

    # PDF
    for i in range(quotas["pdf"]):
        name = _random_basename(rng, "offerta") + ".pdf"
        _write_binary(_work_dir() / name, pdf_bytes + _unique_suffix(i, "pdf"))
        n_created += 1

    # DOCX (zip-based: appendere bytes oltre il central directory non rompe
    # i parser zip moderni, e l'hash dei primi 1MB cambia con prefisso? No:
    # il suffisso non tocca i primi 1MB. Usiamo un campo metadato univoco.
    # Approccio più semplice: scrivi il template + suffisso, e accetta che
    # python-docx potrebbe lamentarsi. Test sul mini-corpus seguente.).
    # In realtà per DOCX/XLSX (zip) un trailing comment NON funziona e può
    # rompere il parser. Allora rigeneriamo il binario con un metadato per-file
    # — più lento ma comunque < pochi secondi.
    import docx as _docx_mod
    import openpyxl as _openpyxl_mod

    for i in range(quotas["docx"]):
        prefix = rng.choice(["contratto", "mail", "verbale", "offerta"])
        name = _random_basename(rng, prefix) + ".docx"
        # Rigenera in-memory un DOCX con paragrafo univoco.
        doc = _docx_mod.Document()
        doc.add_heading("Offerta commerciale", level=1)
        doc.add_paragraph(f"Riferimento univoco: STRESS-DOCX-{i:05d}")
        doc.add_paragraph("Cliente: Bianchi Impianti SpA")
        doc.add_paragraph(
            "Confermiamo disponibilità a fornire le quantità richieste."
        )
        buf = io.BytesIO()
        doc.save(buf)
        _write_binary(_work_dir() / name, buf.getvalue())
        n_created += 1

    for i in range(quotas["xlsx"]):
        prefix = rng.choice(["listino", "anagrafica", "ddt"])
        name = _random_basename(rng, prefix) + ".xlsx"
        wb = _openpyxl_mod.Workbook()
        ws = wb.active
        ws.title = "Listino"
        ws.append(["Codice", "Descrizione", "Prezzo EUR"])
        ws.append([f"STRESS-{i:05d}", "Item sintetico", 18.50 + (i % 50)])
        ws.append(["VL-DN50", "Valvola sfera DN50", 145.00])
        buf = io.BytesIO()
        wb.save(buf)
        _write_binary(_work_dir() / name, buf.getvalue())
        n_created += 1

    # TXT / MD
    for i in range(quotas["txt"]):
        ext = ".md" if rng.random() < 0.4 else ".txt"
        prefix = rng.choice(["note", "appunti", "readme", "todo"])
        name = _random_basename(rng, prefix) + ext
        body = (
            f"Note interne per il file {i}.\n"
            "Cliente: contesto consulenza. Fatturazione, scadenze, riferimenti.\n"
        )
        _write_text(_work_dir() / name, body)
        n_created += 1

    # Noise: file da SKIPPARE (estensioni in _SKIP_EXTENSIONS + node_modules + lock).
    # Distribuzione:
    # - 50% in media/foto (HEIC/JPG)
    # - 20% in media/video (MOV/MP4)
    # - 15% in build/ (lock files, dmg, exe)
    # - 15% in node_modules/ deep tree
    noise_total = quotas["noise"]
    n_foto = int(noise_total * 0.5)
    n_video = int(noise_total * 0.2)
    n_build = int(noise_total * 0.15)
    n_nodemod = noise_total - n_foto - n_video - n_build

    noise_byte = b"\x00" * 64  # 64 byte di placeholder, irrilevanti
    for i in range(n_foto):
        ext = rng.choice([".heic", ".jpg", ".jpeg", ".png"])
        name = f"IMG_{rng.randint(1000, 9999)}_{i}{ext}"
        _write_binary(target_dir / "media" / "foto" / name, noise_byte)
        n_created += 1
    for i in range(n_video):
        ext = rng.choice([".mov", ".mp4"])
        name = f"clip_{i}{ext}"
        _write_binary(target_dir / "media" / "video" / name, noise_byte)
        n_created += 1
    for i in range(n_build):
        ext_pool = [".dmg", ".exe", ".pkg", ".iso", ".lock", ".zip"]
        ext = rng.choice(ext_pool)
        name = f"installer-{i}{ext}"
        _write_binary(target_dir / "build" / name, noise_byte)
        n_created += 1
    for i in range(n_nodemod):
        # File dentro node_modules (skippato come excluded dir component).
        sub = rng.choice(["lodash", "react", "vite", "typescript"])
        name = f"index_{i}.js"
        _write_binary(
            target_dir / "build" / "node_modules" / sub / name, b"// stub\n"
        )
        n_created += 1

    if verbose:
        print(
            f"[generate] creati {n_created} file in {target_dir} "
            f"(target era {total_files})"
        )
    return target_dir


def main(argv: list[str] | None = None) -> int:
    """Entrypoint CLI."""
    parser = argparse.ArgumentParser(
        description=(
            "Genera il corpus sintetico per gli stress test di Custodia "
            "Sprint 2a."
        )
    )
    parser.add_argument(
        "--target",
        type=Path,
        default=Path(tempfile.gettempdir()) / "custodia-stress-corpus",
        help="Cartella di destinazione (default $TMPDIR/custodia-stress-corpus).",
    )
    parser.add_argument(
        "--count",
        type=int,
        default=5000,
        help="Numero di file da generare (default 5000).",
    )
    parser.add_argument(
        "--seed",
        type=int,
        default=42,
        help="Seed RNG per layout deterministico (default 42).",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Stampa progressi su stdout.",
    )
    args = parser.parse_args(argv)
    out = generate_corpus(
        args.target,
        total_files=args.count,
        seed=args.seed,
        verbose=args.verbose,
    )
    print(out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
