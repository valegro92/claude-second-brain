"""
Test dei parser PDF/DOCX/XLSX.

Le fixture sono generate **runtime** usando le stesse librerie che usiamo per
parsare. Questo ci dà file binari portabili, riproducibili e leggibili senza
dover committare blob nel repo.
"""

from __future__ import annotations

import io
from pathlib import Path

import docx as docx_lib
import openpyxl
import pypdf
import pytest
from pypdf import PdfWriter

from custodia_cli.connectors.base import ParserError
from custodia_cli.connectors.parsers import parse_docx, parse_pdf, parse_xlsx


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def _make_pdf_with_text(text_per_page: list[str]) -> bytes:
    """Genera un PDF testuale minimale via pypdf + reportlab-free path.

    Nota: pypdf da solo NON costruisce PDF "ricchi" con testo arbitrario.
    Per evitare dipendenze pesanti (reportlab), generiamo il PDF tramite la API
    a basso livello: una pagina vuota + un content stream con un Tj operator.
    Questo è abbastanza per testare il path estrazione.
    """
    from pypdf.generic import (
        ArrayObject,
        ContentStream,
        DecodedStreamObject,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    writer = PdfWriter()
    for text in text_per_page:
        page = writer.add_blank_page(width=612, height=792)
        # Costruiamo un content stream "BT /F1 12 Tf 72 720 Td (testo) Tj ET"
        escaped = (
            text.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )
        stream_data = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1")
        content = DecodedStreamObject()
        content.set_data(stream_data)
        page[NameObject("/Contents")] = content
        # Aggiungiamo un font Type1 standard (Helvetica) come /F1.
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        resources = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
        )
        page[NameObject("/Resources")] = resources

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_docx(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    """Genera un DOCX in memoria."""
    doc = docx_lib.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    for table in tables or []:
        if not table:
            continue
        n_cols = len(table[0])
        word_table = doc.add_table(rows=len(table), cols=n_cols)
        for r_idx, row in enumerate(table):
            for c_idx, cell in enumerate(row):
                word_table.cell(r_idx, c_idx).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx(sheets: dict[str, list[list[object]]]) -> bytes:
    """Genera un XLSX in memoria. ``sheets`` mappa nome→righe."""
    wb = openpyxl.Workbook()
    # Rimuovi lo sheet di default e ricrea coi nomi voluti.
    default = wb.active
    wb.remove(default)
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Fixture pytest: scrive i file su disco una volta sola.
# ---------------------------------------------------------------------------

FIXTURES_DIR = Path(__file__).parent / "fixtures" / "parsers"


@pytest.fixture(scope="module", autouse=True)
def _ensure_fixtures() -> None:
    """Crea i file di test in tests/fixtures/parsers/ se assenti."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

    pdf_path = FIXTURES_DIR / "sample.pdf"
    if not pdf_path.exists():
        pdf_bytes = _make_pdf_with_text(
            ["Pagina uno con testo Custodia", "Seconda pagina cliente Rossetto", "Pagina finale"]
        )
        pdf_path.write_bytes(pdf_bytes)

    docx_path = FIXTURES_DIR / "sample.docx"
    if not docx_path.exists():
        docx_bytes = _make_docx(
            paragraphs=["Contratto fornitura", "Cliente: Rossetto Laminazioni SRL"],
            tables=[[["Voce", "Valore"], ["Sconto", "5%"]]],
        )
        docx_path.write_bytes(docx_bytes)

    xlsx_path = FIXTURES_DIR / "sample.xlsx"
    if not xlsx_path.exists():
        xlsx_bytes = _make_xlsx(
            {
                "Listino": [["Articolo", "Prezzo"], ["Lamiera 2mm", 42.5], ["Lamiera 3mm", 58.0]],
                "Clienti": [["Nome", "PIVA"], ["Rossetto SRL", "01234567890"]],
            }
        )
        xlsx_path.write_bytes(xlsx_bytes)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def test_parse_pdf_from_path() -> None:
    text = parse_pdf(FIXTURES_DIR / "sample.pdf")
    assert "Custodia" in text
    assert "Rossetto" in text
    assert "finale" in text


def test_parse_pdf_from_bytes() -> None:
    raw = (FIXTURES_DIR / "sample.pdf").read_bytes()
    text = parse_pdf(raw)
    assert "Custodia" in text


def test_parse_pdf_corrupted_raises_parser_error() -> None:
    with pytest.raises(ParserError):
        parse_pdf(b"questo non e' un pdf valido")


def test_parse_pdf_encrypted_returns_empty(tmp_path: Path) -> None:
    """PDF cifrato con password non vuota → ritorna stringa vuota (warning log)."""
    src = _make_pdf_with_text(["secret content"])
    reader = pypdf.PdfReader(io.BytesIO(src))
    writer = pypdf.PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    writer.encrypt(user_password="hunter2", owner_password="hunter2")
    enc_path = tmp_path / "enc.pdf"
    with enc_path.open("wb") as f:
        writer.write(f)

    text = parse_pdf(enc_path)
    assert text == ""


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_parse_docx_from_path() -> None:
    text = parse_docx(FIXTURES_DIR / "sample.docx")
    assert "Contratto fornitura" in text
    assert "Rossetto Laminazioni SRL" in text
    # Tabelle linearizzate con " | ": ancora il contratto del separator.
    assert "Sconto | 5%" in text


def test_parse_docx_from_bytes() -> None:
    raw = (FIXTURES_DIR / "sample.docx").read_bytes()
    text = parse_docx(raw)
    assert "Contratto fornitura" in text


def test_parse_docx_corrupted_raises() -> None:
    with pytest.raises(ParserError):
        parse_docx(b"non e' un docx")


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def test_parse_xlsx_from_path() -> None:
    text = parse_xlsx(FIXTURES_DIR / "sample.xlsx")
    assert "## Sheet: Listino" in text
    assert "## Sheet: Clienti" in text
    assert "Lamiera 2mm" in text
    assert "42.5" in text
    assert "01234567890" in text


def test_parse_xlsx_from_bytes() -> None:
    raw = (FIXTURES_DIR / "sample.xlsx").read_bytes()
    text = parse_xlsx(raw)
    assert "## Sheet: Listino" in text


def test_parse_xlsx_corrupted_raises() -> None:
    with pytest.raises(ParserError):
        parse_xlsx(b"non e' un xlsx")


# ---------------------------------------------------------------------------
# DOCX header/footer (FIX B11)
# ---------------------------------------------------------------------------


def test_parse_docx_extracts_header_and_footer(tmp_path: Path) -> None:
    """Header e footer di un DOCX devono finire nel testo estratto, con
    prefisso ``## Header:`` / ``## Footer:``.

    Caso d'uso reale: offerta italiana con ragione sociale + P.IVA in header.
    """
    doc = docx_lib.Document()
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "ACME SRL — P.IVA 12345"
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Pagina 1 di 1 — riservato"
    doc.add_paragraph("Offerta n. 7")

    target = tmp_path / "with_header.docx"
    doc.save(target)

    text = parse_docx(target)
    assert "ACME SRL" in text
    assert "P.IVA 12345" in text
    assert "Offerta n. 7" in text
    assert "Pagina 1 di 1" in text
    # Verifica i marker:
    assert "## Header:" in text
    assert "## Footer:" in text


# ---------------------------------------------------------------------------
# Zip-bomb guard (FIX B12)
# ---------------------------------------------------------------------------


def test_check_zip_uncompressed_size_rejects_bomb(tmp_path: Path) -> None:
    """Un ZIP la cui somma uncompressed supera la soglia solleva ParserError.

    Usiamo soglia molto piccola per evitare di scrivere centinaia di MB su disco.
    """
    import zipfile

    from custodia_cli.connectors.parsers import check_zip_uncompressed_size

    bomb = tmp_path / "bomb.zip"
    # Scriviamo 10 KB di dati reali; la soglia da 1 KB li rifiuta.
    with zipfile.ZipFile(bomb, "w") as zf:
        zf.writestr("payload.bin", b"x" * 10_000)

    with pytest.raises(ParserError, match="zip-bomb"):
        check_zip_uncompressed_size(bomb, max_uncompressed_bytes=1024)


def test_check_zip_uncompressed_size_allows_normal_file(tmp_path: Path) -> None:
    """Un ZIP sotto soglia passa il check senza errori."""
    import zipfile

    from custodia_cli.connectors.parsers import check_zip_uncompressed_size

    ok = tmp_path / "ok.zip"
    with zipfile.ZipFile(ok, "w") as zf:
        zf.writestr("payload.bin", b"x" * 100)

    # Non solleva (soglia default 200 MB).
    check_zip_uncompressed_size(ok)


def test_check_zip_uncompressed_size_rejects_invalid_zip(tmp_path: Path) -> None:
    """Bytes non-ZIP → ParserError ('Archivio ZIP non valido')."""
    from custodia_cli.connectors.parsers import check_zip_uncompressed_size

    with pytest.raises(ParserError):
        check_zip_uncompressed_size(b"not a zip")


def test_parse_xlsx_rejects_zip_bomb_via_monkeypatch(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """``parse_xlsx`` chiama il size-check al primo passo: lo verifichiamo
    facendo abbassare la soglia."""
    import zipfile

    # Genera un XLSX reale di pochi KB.
    bomb = tmp_path / "small.xlsx"
    import openpyxl

    wb = openpyxl.Workbook()
    wb.active.append(["hi"])
    wb.save(bomb)

    # Forziamo la soglia globale a 10 byte → il file (anche piccolo) la supera.
    from custodia_cli.connectors import parsers as parsers_mod

    monkeypatch.setattr(parsers_mod, "MAX_UNCOMPRESSED_BYTES", 10)

    with pytest.raises(ParserError, match="zip-bomb"):
        parse_xlsx(bomb)
