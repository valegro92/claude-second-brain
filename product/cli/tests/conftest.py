"""
pytest conftest condiviso: genera i file binari del fixture ``finto-drive`` la
prima volta che servono, poi li lascia cached su disco (più piccoli del DOCX/
XLSX se committati e non inquinano il diff della repo).

I file generati sono:
- ``Commerciale 2024/fattura-rossetto-001.pdf``
- ``Commerciale 2024/offerta-bianchi-valvole.docx``
- ``Commerciale 2024/listino-2024.xlsx``

Sono prodotti tramite pypdf/python-docx/openpyxl (già nel requirements di
Custodia per U3). La generazione è idempotente: se il file esiste, non viene
ri-creato. Per forzare il refresh, cancellarli a mano.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest

FIXTURE_ROOT = Path(__file__).parent / "fixtures" / "finto-drive"


def _generate_pdf(target: Path) -> None:
    """Genera un PDF minimale ma valido con pypdf.

    Contiene testo plain ("fattura ... Rossetto Laminazioni SRL ..."), niente
    OCR magic richiesto in lettura.
    """
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    writer = PdfWriter()
    # Pagina A4 vuota; iniettiamo un content stream con BT/ET e un Tj.
    page = writer.add_blank_page(width=595, height=842)
    text_lines = [
        "FATTURA n. 001/2024",
        "Cliente: Rossetto Laminazioni SRL",
        "P.IVA: 03421560289",
        "Data: 12 febbraio 2024",
        "Importo totale: EUR 4.250,00",
        "Descrizione: Fornitura lamiere zincate spess. 1.2mm",
    ]
    # Costruiamo un content stream PDF semplice con font Helvetica.
    # pypdf supporta l'inserimento via ContentStream, ma il modo piu' robusto
    # e' usare un PDF source pre-fatto. Qui usiamo reportlab se disponibile,
    # altrimenti scriviamo manualmente un content stream con un Tj per linea.

    # Aggiungiamo font Helvetica al resource dictionary della pagina.
    font_obj = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font_obj)  # type: ignore[attr-defined]

    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): font_ref}
            ),
        }
    )
    page[NameObject("/Resources")] = resources

    # Content stream: posiziona testo e mostra ogni riga.
    stream_lines = ["BT", "/F1 11 Tf"]
    y = 800
    for line in text_lines:
        # Escape parentesi PDF.
        safe = line.replace("(", r"\(").replace(")", r"\)")
        stream_lines.append(f"1 0 0 1 50 {y} Tm ({safe}) Tj")
        y -= 20
    stream_lines.append("ET")
    content_data = "\n".join(stream_lines).encode("latin-1")

    from pypdf.generic import StreamObject

    stream = StreamObject()
    stream.set_data(content_data)
    stream_ref = writer._add_object(stream)  # type: ignore[attr-defined]
    page[NameObject("/Contents")] = stream_ref

    target.parent.mkdir(parents=True, exist_ok=True)
    with target.open("wb") as fh:
        writer.write(fh)


def _generate_docx(target: Path) -> None:
    """Genera un DOCX finto con python-docx."""
    import docx

    doc = docx.Document()
    doc.add_heading("Offerta commerciale", level=1)
    doc.add_paragraph("Cliente: Bianchi Impianti SpA")
    doc.add_paragraph("Riferimento: OFF-2024-014")
    doc.add_paragraph("Data: 8 marzo 2024")
    doc.add_paragraph(
        "Oggetto: fornitura valvole industriali per nuovo impianto idraulico."
    )
    doc.add_paragraph(
        "Confermiamo disponibilità a fornire n.50 valvole DN50 PN16 al "
        "prezzo unitario di EUR 145,00 + IVA. Consegna 30 gg."
    )
    table = doc.add_table(rows=2, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Codice"
    hdr[1].text = "Descrizione"
    hdr[2].text = "Prezzo"
    row = table.rows[1].cells
    row[0].text = "VL-DN50-PN16"
    row[1].text = "Valvola a sfera DN50 PN16"
    row[2].text = "145,00 EUR"

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def _generate_xlsx(target: Path) -> None:
    """Genera un XLSX finto con openpyxl."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Listino 2024"
    ws.append(["Codice", "Descrizione", "Prezzo EUR"])
    ws.append(["LAM-12-ZN", "Lamiera zincata 1.2mm", 18.50])
    ws.append(["LAM-15-ZN", "Lamiera zincata 1.5mm", 22.00])
    ws.append(["VL-DN50", "Valvola sfera DN50", 145.00])
    ws.append(["VL-DN80", "Valvola sfera DN80", 215.00])

    target.parent.mkdir(parents=True, exist_ok=True)
    wb.save(target)


def _ensure_finto_drive_binaries() -> None:
    """Idempotente: crea i file mancanti del fixture finto-drive.

    Garantisce che TUTTI i file attesi dai test esistano, anche quelli
    plaintext, le immagini placeholder, ``.git/HEAD`` e ``~$temp.docx``.
    Se un file esiste già viene rispettato (nessun overwrite).
    """
    pdf_target = FIXTURE_ROOT / "Commerciale 2024" / "fattura-rossetto-001.pdf"
    docx_target = FIXTURE_ROOT / "Commerciale 2024" / "offerta-bianchi-valvole.docx"
    xlsx_target = FIXTURE_ROOT / "Commerciale 2024" / "listino-2024.xlsx"
    txt_target = FIXTURE_ROOT / "Comunicazioni" / "email-torrelli-2024-03.txt"
    readme_target = FIXTURE_ROOT / "README.md"
    git_head_target = FIXTURE_ROOT / ".git" / "HEAD"
    temp_target = FIXTURE_ROOT / "~$temp.docx"
    img_targets = [
        FIXTURE_ROOT / "archivio" / "foto-prodotti" / "prodotto-1.jpg",
        FIXTURE_ROOT / "archivio" / "foto-prodotti" / "prodotto-2.jpg",
        FIXTURE_ROOT / "archivio" / "foto-prodotti" / "prodotto-3.png",
    ]

    if not pdf_target.exists():
        _generate_pdf(pdf_target)
    if not docx_target.exists():
        _generate_docx(docx_target)
    if not xlsx_target.exists():
        _generate_xlsx(xlsx_target)
    if not txt_target.exists():
        txt_target.parent.mkdir(parents=True, exist_ok=True)
        txt_target.write_text(
            "Da: Torrelli SRL\n"
            "Oggetto: richiesta preventivo lamiere\n"
            "Buongiorno, vorremmo un preventivo per 200 mq di lamiere zincate.\n",
            encoding="utf-8",
        )
    if not readme_target.exists():
        readme_target.write_text(
            "# Finto Drive fixture\n\n"
            "Cartella di test per Custodia. Contiene PDF/DOCX/XLSX/TXT/MD reali.\n",
            encoding="utf-8",
        )
    if not git_head_target.exists():
        git_head_target.parent.mkdir(parents=True, exist_ok=True)
        git_head_target.write_text("ref: refs/heads/main\n", encoding="utf-8")
    if not temp_target.exists():
        temp_target.write_bytes(b"\x00\x01\x02 word-temp-lock placeholder")
    for img in img_targets:
        if not img.exists():
            img.parent.mkdir(parents=True, exist_ok=True)
            img.write_bytes(b"")  # placeholder 0-byte: testa solo lo skip ext


@pytest.fixture(scope="session", autouse=True)
def finto_drive_binaries() -> None:
    """Genera (se necessario) i binari del fixture all'avvio della sessione."""
    _ensure_finto_drive_binaries()


@pytest.fixture
def finto_drive_root() -> Path:
    """Path al fixture finto-drive (binari garantiti presenti)."""
    _ensure_finto_drive_binaries()
    return FIXTURE_ROOT
